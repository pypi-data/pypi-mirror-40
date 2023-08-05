# -*- coding: utf-8 -*-
'''
@author: KaueBonfim
'''
"""Importando bibliotecas externas"""
import pyautogui
import json
import behave2cucumber
import ast
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Length
import os,shutil
import re
"""Importando bibliotecas internas"""
from Pyautomators.Ambiente import _tratar_path

###- annotation:
###    mainTitle: Pyautomators - Modulo Documentação
###    text_description:
###        - "Este modulo tem o intuito de trabalhar a geração de artefatos, trabalhando com a saidas da execução dos testes"
###    ex:
###        - python: "from Pyautomators import Documentacao"

def printarTela(NomeArquivo):
###- annotation:
###    title: Funções
###    text_description:
###        - "**printarTela(NomeArquivo:str)**:"
###    parameters:
###        - NomeArquivo: ":str=Nome do arquivo"
###    text_description1:
###        - "Esta função retira prints da tela e grava em um arquivo."
###    ex:
###        - python: "printarTela('valor.png')"

    """ajustar a path e tirar um print com o nome passado"""
    nome=_tratar_path(NomeArquivo)
    """Tira o print"""
    pyautogui.screenshot(nome) 

def print_local(xi,yi,xf,yf,NomeArquivo):
###- annotation:
###    text_description:
###        - "**print_local(xi:int,yi:int,xf:int,yf:int,NomeArquivo:str)**:"
###    parameters:
###        - xi-yi-xf-yf: ":int(obrigatorio)= Coordenadas iniciais e finais para enquadrar o print."
###        - NomeArquivo: ":str=Nome do arquivo"
###    text_description1:
###        - "Esta função retira prints da tela e grava em um arquivo."
###    ex:
###        - python: "print_local(10,200,100,1000'valor.png')"	

    """faz uma função lambida que retira a largura e a altura"""
    result= lambda a,b:b-a 
    xd=result(xi,xf)
    yd=result(yi,yf)
    """ajustar a path e tirar um print com o nome passado e as medidas iniciais"""
    nome=_tratar_path(NomeArquivo)
    """Tira o print"""
    pyautogui.screenshot(nome,region=(xi,yi, xd, yd))
   
def tranforma_cucumber(NomeArquivo,novo=None):
###- annotation:
###    text_description:
###        - "**tranforma_cucumber(NomeArquivo:str,novo:str)**: "
###    parameters:
###        - NomeArquivo: ":str=Nome do arquivo"
###        - novo: ":str= Nome de um novo arquivo caso necessario gerar os dois, sendo o segundo o padrão json do Cucumber."
###    text_description1:
###        - "Esta função transforma o padrão de report json do Pyautomators, em um padrão compativel com o Cucumber."
###    ex:
###        - python: "tranforma_cucumber('teste.json','teste2.json')"
###        - python: "tranforma_cucumber('teste.json')"

    nome=_tratar_path(NomeArquivo)
    valor=""
    """Abertura do arquivo com os resultados"""
    with open(nome) as behave_json:
        """Transformando o padrão behave em cucumber"""
        cucumber_json = behave2cucumber.convert(json.load(behave_json))
        """Encontrando a duração do teste"""
        for element in cucumber_json:
            elemento=element["elements"]
            for lista in elemento:
                listaa=lista["steps"]
                for lis in listaa:
                    li=lis["result"]["duration"]
                    """Convertendo a quantidade de horas"""
                    lis["result"]["duration"]=int(li*1000000000)
        """Transformando de aspas  simples para aspas duplas."""
        valor=ast.literal_eval(str(cucumber_json).replace("\'",'"""').replace("'",'"'))
    """Caso o nome do segundoo arquivo exita ele grava no segundo arquivo e mantei o primeiro intacto, se não a modificação e salva no primeiro"""
    if(novo is None):
        novo=nome
    """Abrindo o arquivo """
    arquivo = open(novo,'w')  
    """
    Ajustando json"""
    conteudo=json.dumps(valor,indent=4)   
    """Salvando no arquivo"""
    arquivo.write(conteudo)
    """Fechando o arquivo"""
    arquivo.close()
    

def criar_documento_cliente(Nome_documento,json_do_resultado,local_images):
###- annotation:
###    text_description:
###        - "**criar_documento_cliente(Nome_documento,json_do_resultado,local_images)**:"
###    parameters:
###        - Nome_documento: ":str=Nome do PDF de saida para o json"
###        - json_do_resultado: ":str= Caminho do Arquivo json que será formado o resultado"
###        - local_images: ":str= Caminho da pasta aonde tem as imagens dos steps"
###    text_description1:
###        - "Esta função cria um documento PDF com os resultados dos testes executados\n"
###        - "Para vincular as imagens a um teste, o mesmo deve ter como nome da imagem, o nome da funcionalidade,cenario e step a qual a imagem pertence"
###    ex:
###        - python: "criar_documento_cliente('teste.pdf','docs/reports/test.json','docs/')"


        """criando Objeto Documneto, que cria o documento"""
        documento = Document()
        """Criando Titulo"""
        '''Criar Sumario e Capa'''
        documento.add_heading(u'EVIDENCIAS DE TESTE', level=0)
        '''Centralizar Titulo'''
        '''Perguntar o ICONE para a Equipe'''

        imagens=[]
        lista_de_imagens=os.listdir(local_images)
        for arquivos in lista_de_imagens:
            if(str(arquivos).upper().find('.PNG')!=-1):
                imagens.append(arquivos)
                
        arquivo=open(json_do_resultado)
        lista_features=[]
        for arquivo in json.load(arquivo):
            Feature=arquivo['name']
            for interno in arquivo['elements']:
                if(interno['keyword']=='Cenario' or interno['keyword']=='Scenario'):
                    Scenario=interno['name']
                    try:
                        for steps in interno['steps']:
                            lista_features.append((Feature,Scenario,steps['name'],steps['result']['status']))
                    except:
                        pass
                    
        
                
        for step in lista_features:
            
            
                Feature=step[0]
                Scenario=step[1]
                Step=step[2]
                Status=step[3]
                documento.add_page_break()
                documento.add_heading(u'Descrição de Teste', level=0)
            
                texto_feature = documento.add_paragraph().add_run(u'{}----------------------------------------'.format(Feature))
                font_feature = texto_feature.font
                font_feature.size = Pt(14)
                font_feature.underline = True
                font_feature.bold = True
                 
                texto_cenario = documento.add_paragraph().add_run(u'{}'.format(Scenario))
                font_cenario = texto_cenario.font
                font_cenario.size = Pt(13)
                font_cenario.underline = True
                font_cenario.bold = True
                
                texto_de_resultado = documento.add_paragraph().add_run(u'{}'.format(Status))
                font_texto_resultado = texto_de_resultado.font
                font_texto_resultado.size = Pt(12)
                font_texto_resultado.bold = True
                if(Status=='failed'):
                    font_texto_resultado.color.rgb = RGBColor(255,0,0)
                else:
                    font_texto_resultado.color.rgb = RGBColor(0,255,0)
                texto_step = documento.add_paragraph().add_run(u'{}'.format(Step))
                font_step = texto_step.font
                font_step.size = Pt(11)
                font_step.bold = True
                for imagem in imagens:
                    if(str(imagem).find(Feature)!=-1 and str(imagem).find(Scenario)!=-1 and str(imagem).find(Step)!=-1):
                        valor=os.path.join(local_images,imagem) 
                        documento.add_picture(valor, width = Inches(6.0), height = Inches(4.0))
                                       

        documento.save(Nome_documento)

def zipdocs():
    shutil.make_archive(["docs/docs"], "zip",  base_dir="docs")
    shutil.make_archive('docs/logs',"zip",base_dir="log")